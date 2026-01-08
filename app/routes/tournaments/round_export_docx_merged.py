# app/routes/tournaments/round_export_docx_merged.py
from __future__ import annotations

from dataclasses import dataclass
import io
import zipfile

from flask import flash, redirect, send_file, url_for

from ... import db
from . import bp
from .helpers import _get_tournament


# -----------------------------------------------------------------------------
# Data structures
# -----------------------------------------------------------------------------
@dataclass(frozen=True)
class SeatInfo:
    seat: str            # "A" | "B" | "C" | "D"
    player_no: int       # Teilnehmernummer
    display_name: str    # Anzeigename
    email: str | None    # E-Mail (kann fehlen)


@dataclass(frozen=True)
class TableInfo:
    table_no: int
    seats: dict[str, SeatInfo]  # keys: "A","B","C","D"


# -----------------------------------------------------------------------------
# Helpers: fetch table data
# -----------------------------------------------------------------------------
def _fetch_round_tables(con, tournament_id: int, round_no: int) -> list[TableInfo]:
    rows = db.q(
        con,
        """
        SELECT
          s.table_no,
          s.seat,
          tp.player_no,
          tp.display_name,
          a.nachname,
          a.vorname,
          a.wohnort,
          a.email
        FROM tournament_seats s
        JOIN tournament_participants tp ON tp.id = s.tp_id
        JOIN addresses a ON a.id = tp.address_id
        WHERE s.tournament_id = ? AND s.round_no = ?
        ORDER BY
          s.table_no ASC,
          CASE s.seat
            WHEN 'A' THEN 1
            WHEN 'B' THEN 2
            WHEN 'C' THEN 3
            WHEN 'D' THEN 4
            ELSE 9
          END
        """,
        (int(tournament_id), int(round_no)),
    )

    by_table: dict[int, dict[str, SeatInfo]] = {}
    for r in rows:
        tno = int(r["table_no"])
        seat = str(r["seat"] or "").strip().upper()

        disp = (r["display_name"] or "").strip()
        if not disp:
            nn = (r["nachname"] or "").strip()
            vn = (r["vorname"] or "").strip()
            wo = (r["wohnort"] or "").strip()
            disp = f"{nn}, {vn}"
            if wo:
                disp += f" · {wo}"

        si = SeatInfo(
            seat=seat,
            player_no=int(r["player_no"] or 0),
            display_name=disp,
            email=(str(r["email"]).strip() if r["email"] else None),
        )
        by_table.setdefault(tno, {})
        by_table[tno][seat] = si

    out: list[TableInfo] = []
    for tno in sorted(by_table.keys()):
        seats = by_table[tno]
        if not all(k in seats for k in ("A", "B", "C", "D")):
            continue
        out.append(TableInfo(table_no=tno, seats=seats))
    return out


def _fetch_single_table(con, tournament_id: int, round_no: int, table_no: int) -> TableInfo | None:
    tables = _fetch_round_tables(con, tournament_id, round_no)
    for t in tables:
        if int(t.table_no) == int(table_no):
            return t
    return None


def _safe_filename(s: str) -> str:
    s2 = "".join(ch for ch in (s or "") if ch.isalnum() or ch in (" ", "-", "_")).strip()
    s2 = s2.replace(" ", "_")
    return s2 or "Turnier"


# -----------------------------------------------------------------------------
# DOCX builder (no template)
# -----------------------------------------------------------------------------
def _build_merged_docx(*, tournament_title: str, round_no: int, tables: list[TableInfo]) -> bytes:
    """
    Baut ein DOCX, das für jeden Tisch zwei Seiten enthält (Seite 1/2 und 2/2).
    Diese Funktion enthält den kompletten Layout-/Tabellenbau.
    """
    from pathlib import Path

    from flask import current_app
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Mm, Pt, RGBColor

    doc = Document()

    # --- Page setup ---
    sec = doc.sections[0]
    sec.top_margin = Mm(10)
    sec.bottom_margin = Mm(10)
    sec.left_margin = Mm(10)
    sec.right_margin = Mm(10)

    # header/footer distances + clear content (no header/footer at all)
    try:
        sec.header_distance = Mm(0)
        sec.footer_distance = Mm(0)
    except Exception:
        pass
    try:
        sec.different_first_page_header_footer = False
    except Exception:
        pass
    try:
        sec.header.is_linked_to_previous = False
        sec.footer.is_linked_to_previous = False
    except Exception:
        pass
    for p in sec.header.paragraphs:
        p.text = ""
    for p in sec.footer.paragraphs:
        p.text = ""

    # --- Global default font: Source Sans Pro, 12pt ---
    style = doc.styles["Normal"]
    style.font.name = "Source Sans Pro"
    style.font.size = Pt(12)
    try:
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
    except Exception:
        pass

    # Column widths (9 columns)
    w0 = Cm(1.6)
    w = Cm(2.2)
    col_widths = [w0] + [w] * 8

    # Shades
    SHADE_HEADER = "F2F2F2"
    SHADE_LIGHT = "FAFAFA"
    SHADE_SUM = "F7F7F7"
    SHADE_BOTTOM = "F2F2F2"
    SHADE_WHITE = "FFFFFF"

    # -----------------------------------------------------------------------------
    # Dealer marker (Geber-Markierung): app/static/branding/favicon.png
    # -----------------------------------------------------------------------------
    dealer_icon_path = Path(current_app.root_path) / "static" / "branding" / "favicon.png"
    DEALER_SEATS = ("A", "B", "C", "D")
    DEALER_COL_BY_SEAT = {"A": 1, "B": 3, "C": 5, "D": 7}

    def _dealer_seat_for_game(game_no: int) -> str:
        return DEALER_SEATS[(int(game_no) - 1) % 4]

    def _add_dealer_marker_to_cell(cell) -> None:
        if not dealer_icon_path.exists():
            return
        cell.text = ""
        if not cell.paragraphs:
            cell.add_paragraph()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        except Exception:
            pass
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        run = p.add_run()
        try:
            run.add_picture(str(dealer_icon_path), width=Mm(4))
        except Exception:
            run.add_picture(str(dealer_icon_path))

    # -------- low-level XML helpers --------
    def _rgb(color: str | None) -> RGBColor | None:
        if not color:
            return None
        c = str(color).strip().lstrip("#")
        if len(c) == 3:
            c = "".join(ch * 2 for ch in c)
        if len(c) != 6:
            return None
        try:
            return RGBColor.from_string(c.upper())
        except Exception:
            return None

    def _set_cell_shading(cell, *, fill: str) -> None:
        c = str(fill or "").strip().lstrip("#")
        if len(c) == 3:
            c = "".join(ch * 2 for ch in c)
        if len(c) != 6:
            return
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), c.upper())

    def _set_table_cell_margins(
        tbl,
        *,
        top_tw: int = 60,
        bottom_tw: int = 60,
        left_tw: int = 90,
        right_tw: int = 90,
    ) -> None:
        tblPr = tbl._tbl.tblPr
        mar = tblPr.find(qn("w:tblCellMar"))
        if mar is None:
            mar = OxmlElement("w:tblCellMar")
            tblPr.append(mar)

        def _set(tag: str, val: int):
            el = mar.find(qn(f"w:{tag}"))
            if el is None:
                el = OxmlElement(f"w:{tag}")
                mar.append(el)
            el.set(qn("w:w"), str(int(val)))
            el.set(qn("w:type"), "dxa")

        _set("top", top_tw)
        _set("bottom", bottom_tw)
        _set("left", left_tw)
        _set("right", right_tw)

    def _set_table_borders(tbl, *, outer_pt: float = 2.0, inner_pt: float = 1.0) -> None:
        outer_sz = str(int(round(outer_pt * 8)))
        inner_sz = str(int(round(inner_pt * 8)))

        tblPr = tbl._tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tblPr.append(borders)

        def _border(tag: str, sz: str):
            el = borders.find(qn(f"w:{tag}"))
            if el is None:
                el = OxmlElement(f"w:{tag}")
                borders.append(el)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")

        _border("top", outer_sz)
        _border("left", outer_sz)
        _border("bottom", outer_sz)
        _border("right", outer_sz)
        _border("insideH", inner_sz)
        _border("insideV", inner_sz)

    def _set_cell_borders(cell, *, left=None, right=None, top=None, bottom=None) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)

        def _apply(side: str, spec):
            if spec is None:
                return
            el = tcBorders.find(qn(f"w:{side}"))
            if el is None:
                el = OxmlElement(f"w:{side}")
                tcBorders.append(el)
            el.set(qn("w:val"), spec.get("val", "single"))
            el.set(qn("w:sz"), str(int(spec.get("sz", 8))))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), spec.get("color", "000000"))

        _apply("left", left)
        _apply("right", right)
        _apply("top", top)
        _apply("bottom", bottom)

    # -------- text helpers --------
    def _fit_text_lines(text: str, *, max_chars_per_line: int, max_lines: int) -> list[str]:
        t = " ".join((text or "").split())
        if not t:
            return [""]
        words = t.split(" ")
        lines: list[str] = []
        cur = ""
        for w in words:
            cand = w if not cur else (cur + " " + w)
            if len(cand) <= max_chars_per_line:
                cur = cand
            else:
                if cur:
                    lines.append(cur)
                cur = w
                if len(lines) >= max_lines:
                    break
        if len(lines) < max_lines and cur:
            lines.append(cur)

        joined = " ".join(lines)
        if len(joined) < len(t):
            last = lines[-1]
            if len(last) >= max_chars_per_line:
                last = last[: max(0, max_chars_per_line - 1)]
            lines[-1] = last.rstrip(".") + "…"
        return lines[:max_lines]

    def _auto_header_sizes(name: str, email: str) -> tuple[int, int]:
        n = len((name or "").strip())
        e = len((email or "").strip())
        name_pt = 11
        email_pt = 9
        if n > 42:
            name_pt = 10
        if n > 60:
            name_pt = 9
        if e > 32:
            email_pt = 8
        if e > 45:
            email_pt = 7
        return name_pt, email_pt

    def _set_cell_paragraph(
        cell,
        text: str,
        *,
        align,
        pt: int = 12,
        bold: bool = False,
        italic: bool = False,
        color: str = "",
    ) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        try:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        except Exception:
            pass
        run = p.add_run(text)
        run.font.name = "Source Sans Pro"
        run.font.size = Pt(pt)
        run.bold = bool(bold)
        run.italic = bool(italic)
        rgb = _rgb(color)
        if rgb is not None:
            run.font.color.rgb = rgb

    def _clear_cell(cell) -> None:
        cell.text = ""
        if not cell.paragraphs:
            cell.add_paragraph()

    def _set_header_cell(
        cell,
        *,
        seat: str,
        player_no: int,
        name: str,
        email: str,
        include_email: bool,
    ) -> None:
        _clear_cell(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        name_pt, email_pt = _auto_header_sizes(name, email)
        name_lines = _fit_text_lines(name, max_chars_per_line=28, max_lines=2)
        email_lines = _fit_text_lines(email, max_chars_per_line=26, max_lines=2)

        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(0)
            p1.paragraph_format.line_spacing = 1.0
        except Exception:
            pass

        r_pl = p1.add_run(f"Platz {seat}")
        r_pl.bold = True
        r_pl.font.name = "Source Sans Pro"
        r_pl.font.size = Pt(12)

        r_tail = p1.add_run(f" — Teiln.-Nr. {int(player_no)}")
        r_tail.bold = False
        r_tail.font.name = "Source Sans Pro"
        r_tail.font.size = Pt(10)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.line_spacing = 1.0
        except Exception:
            pass
        r2 = p2.add_run("\n".join(name_lines).strip())
        r2.font.name = "Source Sans Pro"
        r2.font.size = Pt(name_pt)

        if not include_email:
            return

        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(0)
            p3.paragraph_format.line_spacing = 1.0
        except Exception:
            pass

        r3a = p3.add_run("E-Mail: ")
        r3a.font.name = "Source Sans Pro"
        r3a.font.size = Pt(email_pt)
        rgb = _rgb("666666")
        if rgb is not None:
            r3a.font.color.rgb = rgb

        r3b = p3.add_run("\n".join(email_lines).strip())
        r3b.font.name = "Source Sans Pro"
        r3b.font.size = Pt(email_pt)
        rgb2 = _rgb("444444")
        if rgb2 is not None:
            r3b.font.color.rgb = rgb2

    def _shade_row(row, *, fill: str) -> None:
        for c in row.cells:
            _set_cell_shading(c, fill=fill)

    def _merge_pairs(row) -> None:
        row.cells[1].merge(row.cells[2])
        row.cells[3].merge(row.cells[4])
        row.cells[5].merge(row.cells[6])
        row.cells[7].merge(row.cells[8])

    def _apply_dashed_internal_separators(tbl, *, row_indices: list[int]) -> None:
        dashed = {"val": "dashed", "sz": 4, "color": "000000"}  # 0.5pt -> 4
        boundaries = [(1, 2), (3, 4), (5, 6), (7, 8)]
        for rix in row_indices:
            row = tbl.rows[rix]
            for a, b in boundaries:
                try:
                    _set_cell_borders(row.cells[a], right=dashed)
                    _set_cell_borders(row.cells[b], left=dashed)
                except Exception:
                    pass

    def _enforce_outer_right_border(tbl, *, outer_pt: float = 2.0) -> None:
        outer_sz = int(round(outer_pt * 8))  # 2pt -> 16
        spec = {"val": "single", "sz": outer_sz, "color": "000000"}
        for row in tbl.rows:
            try:
                _set_cell_borders(row.cells[-1], right=spec)
            except Exception:
                pass

    def _apply_table_layout(tbl, *, page_no: int) -> None:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        for row in tbl.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = col_widths[ci]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        _set_table_cell_margins(tbl, top_tw=30, bottom_tw=30, left_tw=45, right_tw=45)
        _set_table_borders(tbl, outer_pt=2.0, inner_pt=1.0)

        header_h_page1 = 28.0
        header_h_page2 = 16.0

        plusminus_h = 5.5
        carry_h_page2 = 8.3
        game_h = 8.3
        sum_h = 8.3
        bottom_h = 8.3

        heights_mm: dict[int, float] = {}

        if page_no == 1:
            heights_mm[0] = header_h_page1
            heights_mm[1] = plusminus_h
            for rix in range(2, 22):
                heights_mm[rix] = game_h
            heights_mm[22] = sum_h
            heights_mm[23] = bottom_h
            dashed_rows = list(range(2, 22)) + [22]
        else:
            heights_mm[0] = header_h_page2
            heights_mm[1] = carry_h_page2
            heights_mm[2] = plusminus_h
            for rix in range(3, 23):
                heights_mm[rix] = game_h
            heights_mm[rix + 1] = game_h  # no-op safety
            heights_mm[23] = sum_h
            heights_mm[24] = bottom_h
            dashed_rows = list(range(3, 23)) + [23]

        for rix, row in enumerate(tbl.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Mm(heights_mm.get(rix, game_h))

        _apply_dashed_internal_separators(tbl, row_indices=dashed_rows)
        _enforce_outer_right_border(tbl, outer_pt=2.0)

        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    try:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.0
                    except Exception:
                        pass
                    for run in p.runs:
                        run.font.name = "Source Sans Pro"

    def _fill_header_row(tbl, table: TableInfo, *, include_email: bool) -> None:
        _merge_pairs(tbl.rows[0])
        _shade_row(tbl.rows[0], fill=SHADE_HEADER)

        seat_order = ["A", "B", "C", "D"]
        start_cells = [tbl.rows[0].cells[1], tbl.rows[0].cells[3], tbl.rows[0].cells[5], tbl.rows[0].cells[7]]

        for seat, cell in zip(seat_order, start_cells):
            s = table.seats[seat]
            _set_header_cell(
                cell,
                seat=seat,
                player_no=int(s.player_no),
                name=str(s.display_name or ""),
                email=str(s.email or ""),
                include_email=include_email,
            )

    def _fill_carry_in_row(tbl, *, row_index: int) -> None:
        row = tbl.rows[row_index]
        _merge_pairs(row)
        _shade_row(row, fill=SHADE_LIGHT)

        c0 = row.cells[0]
        c0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_paragraph(c0, "Übertrag", align=WD_ALIGN_PARAGRAPH.LEFT, pt=10, color="444444")

        for c in [row.cells[1], row.cells[3], row.cells[5], row.cells[7]]:
            c.text = ""
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _fill_plusminus_row(tbl, *, row_index: int) -> None:
        row = tbl.rows[row_index]
        _merge_pairs(row)
        _shade_row(row, fill=SHADE_LIGHT)

        cells = [row.cells[1], row.cells[3], row.cells[5], row.cells[7]]
        for c in cells:
            c.text = ""
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
            except Exception:
                pass
            run = p.add_run("Plus  + / -  Minus")
            run.italic = True
            run.font.name = "Source Sans Pro"
            run.font.size = Pt(10)
            rgb = _rgb("666666")
            if rgb is not None:
                run.font.color.rgb = rgb

    def _fill_games_and_totals(
        tbl,
        *,
        start_row: int,
        start_game_no: int,
        sum_row: int,
        bottom_row: int,
        bottom_label: str,
        bottom_label_pt: int,
    ) -> None:
        for i in range(20):
            rix = start_row + i
            game_no = start_game_no + i

            _shade_row(tbl.rows[rix], fill=("FCFCFC" if i % 2 == 1 else SHADE_WHITE))

            tbl.rows[rix].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_paragraph(tbl.rows[rix].cells[0], str(game_no), align=WD_ALIGN_PARAGRAPH.RIGHT, pt=12)

            for ci in range(1, 9):
                c = tbl.rows[rix].cells[ci]
                c.text = ""
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            seat = _dealer_seat_for_game(game_no)
            col = DEALER_COL_BY_SEAT[seat]
            _add_dealer_marker_to_cell(tbl.rows[rix].cells[col])

        sumr = tbl.rows[sum_row]
        _shade_row(sumr, fill=SHADE_SUM)
        _set_cell_paragraph(sumr.cells[0], "Summe", align=WD_ALIGN_PARAGRAPH.RIGHT, pt=12, bold=True)
        for ci in range(1, 9):
            sumr.cells[ci].text = ""
            sumr.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        bot = tbl.rows[bottom_row]
        _shade_row(bot, fill=SHADE_BOTTOM)
        _set_cell_paragraph(
            bot.cells[0],
            bottom_label,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            pt=bottom_label_pt,
            bold=True if bottom_label == "Gesamt" else False,
            color="444444" if bottom_label != "Gesamt" else "",
        )
        _merge_pairs(bot)
        for c in [bot.cells[1], bot.cells[3], bot.cells[5], bot.cells[7]]:
            c.text = ""
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_tablesheet_page(table: TableInfo, *, page_no: int) -> None:
        p = doc.add_paragraph()
        try:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.0
        except Exception:
            pass

        right_pos = sec.page_width - sec.left_margin - sec.right_margin
        try:
            ts = p.paragraph_format.tab_stops
            try:
                ts.clear_all()
            except Exception:
                pass
            ts.add_tab_stop(right_pos, alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.SPACES)
        except Exception:
            pass

        r1 = p.add_run(f"{tournament_title}  Runde {int(round_no)}  ")
        r1.font.name = "Source Sans Pro"
        r1.font.size = Pt(16)
        r1.bold = True

        r2 = p.add_run(f"Tisch {int(table.table_no)}")
        r2.font.name = "Source Sans Pro"
        r2.font.size = Pt(24)
        r2.bold = True

        r3 = p.add_run(f"\tSeite {page_no}/2")
        r3.font.name = "Source Sans Pro"
        r3.font.size = Pt(12)
        r3.bold = False

        if page_no == 1:
            tbl = doc.add_table(rows=24, cols=9)
            _apply_table_layout(tbl, page_no=1)
            _fill_header_row(tbl, table, include_email=True)
            _fill_plusminus_row(tbl, row_index=1)
            _fill_games_and_totals(
                tbl,
                start_row=2,
                start_game_no=1,
                sum_row=22,
                bottom_row=23,
                bottom_label="Übertrag",
                bottom_label_pt=10,
            )
        else:
            tbl = doc.add_table(rows=25, cols=9)
            _apply_table_layout(tbl, page_no=2)
            _fill_header_row(tbl, table, include_email=False)
            _fill_carry_in_row(tbl, row_index=1)
            _fill_plusminus_row(tbl, row_index=2)
            _fill_games_and_totals(
                tbl,
                start_row=3,
                start_game_no=21,
                sum_row=23,
                bottom_row=24,
                bottom_label="Gesamt",
                bottom_label_pt=12,
            )

    for idx, table in enumerate(tables):
        _add_tablesheet_page(table, page_no=1)
        doc.add_page_break()
        _add_tablesheet_page(table, page_no=2)
        if idx < len(tables) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_single_docx(*, tournament_title: str, round_no: int, table: TableInfo) -> bytes:
    # Single = merged mit genau einem Tisch
    return _build_merged_docx(tournament_title=tournament_title, round_no=round_no, tables=[table])


# -----------------------------------------------------------------------------
# Route: merged tablesheets DOCX (all tables in one doc)
# Endpoint used in template: tournaments.tournament_round_tablesheets_docx_merged
# -----------------------------------------------------------------------------
@bp.get("/tournaments/<int:tournament_id>/rounds/<int:round_no>/tablesheets-docx-merged")
def tournament_round_tablesheets_docx_merged(tournament_id: int, round_no: int):
    try:
        import docx  # noqa: F401
    except ModuleNotFoundError:
        flash("DOCX-Export nicht verfügbar: Paket 'python-docx' ist nicht installiert.", "error")
        return redirect(url_for("tournaments.tournament_round_view", tournament_id=tournament_id, round_no=round_no))

    with db.connect() as con:
        t = _get_tournament(con, tournament_id)
        if not t:
            flash("Turnier nicht gefunden.", "error")
            return redirect(url_for("tournaments.tournaments_list"))

        tables = _fetch_round_tables(con, tournament_id, round_no)
        if not tables:
            flash(f"Keine vollständigen 4er-Tische für Runde {round_no} gefunden (Auslosung fehlt?).", "error")
            return redirect(url_for("tournaments.tournament_round_view", tournament_id=tournament_id, round_no=round_no))

        title = str(t["title"] or "").strip() or "Turnier"

    payload = _build_merged_docx(tournament_title=title, round_no=int(round_no), tables=tables)

    fn = f"{_safe_filename(title)}_R{int(round_no):02d}_Tische.docx"
    return send_file(
        io.BytesIO(payload),
        as_attachment=True,
        download_name=fn,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        max_age=0,
    )


# -----------------------------------------------------------------------------
# Route: single table tablesheet DOCX
# Endpoint used in results template: tournaments.tournament_round_tablesheet_docx_single
# -----------------------------------------------------------------------------
@bp.get("/tournaments/<int:tournament_id>/rounds/<int:round_no>/tablesheets-docx/<int:table_no>")
def tournament_round_tablesheet_docx_single(tournament_id: int, round_no: int, table_no: int):
    try:
        import docx  # noqa: F401
    except ModuleNotFoundError:
        flash("DOCX-Export nicht verfügbar: Paket 'python-docx' ist nicht installiert.", "error")
        return redirect(url_for("tournaments.tournament_round_view", tournament_id=tournament_id, round_no=round_no))

    with db.connect() as con:
        t = _get_tournament(con, tournament_id)
        if not t:
            flash("Turnier nicht gefunden.", "error")
            return redirect(url_for("tournaments.tournaments_list"))

        table = _fetch_single_table(con, tournament_id, round_no, table_no)
        if not table:
            flash(f"Tisch {table_no} in Runde {round_no} nicht gefunden (oder nicht vollständig).", "error")
            return redirect(url_for("tournaments.tournament_round_view", tournament_id=tournament_id, round_no=round_no))

        title = str(t["title"] or "").strip() or "Turnier"

    payload = _build_single_docx(tournament_title=title, round_no=int(round_no), table=table)
    fn = f"{_safe_filename(title)}_R{int(round_no):02d}_T{int(table_no):02d}.docx"
    return send_file(
        io.BytesIO(payload),
        as_attachment=True,
        download_name=fn,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        max_age=0,
    )


# -----------------------------------------------------------------------------
# Route: all tables as individual DOCX files in a ZIP
# Endpoint used in round template: tournaments.tournament_round_tablesheets_docx_zip
# -----------------------------------------------------------------------------
@bp.get("/tournaments/<int:tournament_id>/rounds/<int:round_no>/tablesheets-docx-zip")
def tournament_round_tablesheets_docx_zip(tournament_id: int, round_no: int):
    try:
        import docx  # noqa: F401
    except ModuleNotFoundError:
        flash("DOCX-Export nicht verfügbar: Paket 'python-docx' ist nicht installiert.", "error")
        return redirect(url_for("tournaments.tournament_round_view", tournament_id=tournament_id, round_no=round_no))

    with db.connect() as con:
        t = _get_tournament(con, tournament_id)
        if not t:
            flash("Turnier nicht gefunden.", "error")
            return redirect(url_for("tournaments.tournaments_list"))

        tables = _fetch_round_tables(con, tournament_id, round_no)
        if not tables:
            flash(f"Keine vollständigen 4er-Tische für Runde {round_no} gefunden (Auslosung fehlt?).", "error")
            return redirect(url_for("tournaments.tournament_round_view", tournament_id=tournament_id, round_no=round_no))

        title = str(t["title"] or "").strip() or "Turnier"

    safe = _safe_filename(title)
    zip_name = f"{safe}_R{int(round_no):02d}_Tischblaetter.zip"

    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for table in tables:
            docx_payload = _build_single_docx(tournament_title=title, round_no=int(round_no), table=table)
            docx_name = f"{safe}_R{int(round_no):02d}_T{int(table.table_no):02d}.docx"
            zf.writestr(docx_name, docx_payload)

    zbuf.seek(0)
    return send_file(
        zbuf,
        as_attachment=True,
        download_name=zip_name,
        mimetype="application/zip",
        max_age=0,
    )