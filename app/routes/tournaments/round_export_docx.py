# app/routes/tournaments/round_export_docx.py
from __future__ import annotations

from dataclasses import dataclass
import io

from flask import flash, redirect, send_file, url_for

from ... import db
from . import bp
from .helpers import _get_tournament


# -----------------------------------------------------------------------------
# Data structures
# -----------------------------------------------------------------------------
@dataclass(frozen=True)
class SeatInfo:
    seat: str            # "A" | "B" | "C" | "D"
    player_no: int       # Teilnehmernummer
    display_name: str    # Anzeigename (oder "Nachname, Vorname · Wohnort")
    email: str | None    # E-Mail (kann fehlen)


@dataclass(frozen=True)
class TableInfo:
    table_no: int
    seats: dict[str, SeatInfo]  # keys: "A","B","C","D"


# -----------------------------------------------------------------------------
# Helpers: fetch table data
# -----------------------------------------------------------------------------
def _fetch_round_tables(con, tournament_id: int, round_no: int) -> list[TableInfo]:
    """
    Liefert pro Tisch eine TableInfo-Struktur inkl. A-D Sitzen.
    Erwartet vorhandene Auslosung in tournament_seats.
    Sortierung: table_no ASC, seat A-D.
    """
    rows = db.q(
        con,
        """
        SELECT
          s.table_no,
          s.seat,
          tp.player_no,
          tp.display_name,
          a.nachname,
          a.vorname,
          a.wohnort,
          a.email
        FROM tournament_seats s
        JOIN tournament_participants tp ON tp.id = s.tp_id
        JOIN addresses a ON a.id = tp.address_id
        WHERE s.tournament_id = ? AND s.round_no = ?
        ORDER BY
          s.table_no ASC,
          CASE s.seat
            WHEN 'A' THEN 1
            WHEN 'B' THEN 2
            WHEN 'C' THEN 3
            WHEN 'D' THEN 4
            ELSE 9
          END
        """,
        (int(tournament_id), int(round_no)),
    )

    by_table: dict[int, dict[str, SeatInfo]] = {}

    for r in rows:
        tno = int(r["table_no"])
        seat = str(r["seat"] or "").strip().upper()

        disp = (r["display_name"] or "").strip()
        if not disp:
            nn = (r["nachname"] or "").strip()
            vn = (r["vorname"] or "").strip()
            wo = (r["wohnort"] or "").strip()
            disp = f"{nn}, {vn}"
            if wo:
                disp += f" · {wo}"

        si = SeatInfo(
            seat=seat,
            player_no=int(r["player_no"] or 0),
            display_name=disp,
            email=(str(r["email"]).strip() if r["email"] else None),
        )

        by_table.setdefault(tno, {})
        by_table[tno][seat] = si

    out: list[TableInfo] = []
    for tno in sorted(by_table.keys()):
        seats = by_table[tno]
        if not all(k in seats for k in ("A", "B", "C", "D")):
            # nur vollständige 4er-Tische exportieren
            continue
        out.append(TableInfo(table_no=tno, seats=seats))

    return out


# -----------------------------------------------------------------------------
# DOCX builder (Variante B: alles per Code)
# -----------------------------------------------------------------------------
def _set_cell_text(cell, text: str, *, align, font_size_pt: int) -> None:
    """
    Setzt Zelltext und formatiert alle Absätze (Ausrichtung + Fontsize).
    """
    # cell.text überschreibt Inhalt und erzeugt i.d.R. einen Absatz
    cell.text = text or ""
    for p in cell.paragraphs:
        p.alignment = align
        for run in p.runs:
            run.font.size = Pt(font_size_pt)


def _clear_cell(cell) -> None:
    """
    Entfernt Inhalt robust.
    """
    cell.text = ""
    # Word lässt gern leere Paragraphen stehen; ist ok


def _set_cell_paragraphs_multiline(cell, lines: list[str], *, align, font_size_pt: int) -> None:
    """
    Schreibt mehrere Zeilen als separate Absätze (sauberer als \n in cell.text).
    """
    # Inhalt leeren
    cell.text = ""
    # erster Absatz existiert bereits:
    p0 = cell.paragraphs[0]
    p0.alignment = align
    p0.clear()  # type: ignore[attr-defined]  # python-docx hat clear() nicht offiziell überall

    # "clear()" ist nicht in allen python-docx Versionen vorhanden -> fallback:
    try:
        # falls clear() nicht existiert
        pass
    except Exception:
        pass

    # Wir machen es sicher: komplett neu befüllen über add_paragraph
    # (erstes paragraph-Objekt bleibt leer, ist aber unschädlich)
    for i, line in enumerate(lines):
        p = cell.add_paragraph(line or "")
        p.alignment = align
        for run in p.runs:
            run.font.size = Pt(font_size_pt)

    # falls ganz am Anfang ein leerer Absatz steht: entfernen geht nicht sauber,
    # aber wir können ihn leer lassen (optisch i.d.R. unauffällig).
    # Alternativ könnte man XML entfernen – lassen wir bewusst stabil.


def _build_one_page(
    doc,
    *,
    tournament_title: str,
    round_no: int,
    table_no: int,
    seats: dict[str, SeatInfo],
    game_start_no: int,
) -> None:
    """
    Baut exakt eine Seite:
      Zeile 1: Header (14pt): Turniername Runde Tisch
      Zeile 2: Tabelle 9 Spalten / 24 Zeilen gemäß Spezifikation
    """
    # lazy imports (nur wenn Route aufgerufen wird)
    # (oben wäre auch ok, aber so bleibt Datei "leicht")
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

    # -------------------------
    # Header-Zeile (14 pt)
    # -------------------------
    header_text = f"{tournament_title}  Runde {int(round_no)}  Tisch {int(table_no)}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(header_text)
    r.font.size = Pt(14)

    # -------------------------
    # Tabelle 9 x 24
    # -------------------------
    rows = 24
    cols = 9
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # feste Spaltenbreiten (cm)
    col_widths = [Cm(1.6)] + [Cm(2.2)] * 8

    # Globale Zellformatierung: vertikal zentrieren + Schrift 12
    for ri in range(rows):
        for ci in range(cols):
            cell = tbl.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Absatz-Ausrichtung:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(12)

    # Spaltenbreite setzen (pro Zelle, weil Word Tabellenbreiten gern pro Zelle speichert)
    for ci, w in enumerate(col_widths):
        for ri in range(rows):
            tbl.cell(ri, ci).width = w

    # Spalte 1 rechtsbündig (Nummerierung / Summe / Übertrag)
    for ri in range(rows):
        cell = tbl.cell(ri, 0)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.size = Pt(12)

    # -------------------------
    # Zeile 1: Header-Merges
    # Merge (2&3), (4&5), (6&7), (8&9) -> 0-index: (1,2),(3,4),(5,6),(7,8)
    # -------------------------
    def merge_pair(row_idx: int, left_col: int) -> None:
        _ = tbl.cell(row_idx, left_col).merge(tbl.cell(row_idx, left_col + 1))

    for left in (1, 3, 5, 7):
        merge_pair(0, left)

    # Inhalte Zeile 1
    # Wir schreiben "Platz X" + darunter Nr, Name, E-Mail (12pt, zentriert)
    def seat_block(letter: str) -> list[str]:
        s = seats[letter]
        return [
            f"Platz {letter}",
            f"Nr.: {int(s.player_no)}",
            str(s.display_name or ""),
            str(s.email or ""),
        ]

    # In gemergten Zellen steht Text (wir nutzen \n, das ist in Word ok)
    # allerdings sieht "separate Absätze" oft besser aus -> wir machen Absätze.
    # Zielzellen: col 1,3,5,7
    seat_cols = {"A": 1, "B": 3, "C": 5, "D": 7}
    for letter, col in seat_cols.items():
        cell = tbl.cell(0, col)
        cell.text = ""  # leeren
        # Wir füllen als Absätze
        # (erste leere paragraph bleibt ggf. stehen, optisch ok; wir können ihn versuchen zu leeren)
        # Wir setzen dann 4 Absätze:
        for line in seat_block(letter):
            pp = cell.add_paragraph(line)
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in pp.runs:
                run.font.size = Pt(12)

    # -------------------------
    # Zeile 2: gleiche Merges + "Plus+ / -Minus"
    # -------------------------
    for left in (1, 3, 5, 7):
        merge_pair(1, left)

    for col in (1, 3, 5, 7):
        c = tbl.cell(1, col)
        c.text = "Plus+ / -Minus"
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(12)

    # -------------------------
    # Zeilen 3-22: Spielnummern (20 Stück)
    # -------------------------
    # row index 2..21
    for i in range(20):
        row_idx = 2 + i
        game_no = game_start_no + i
        c0 = tbl.cell(row_idx, 0)
        c0.text = str(game_no)
        for para in c0.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.size = Pt(12)

    # -------------------------
    # Zeile 23: "Summe" in Spalte 1
    # -------------------------
    tbl.cell(22, 0).text = "Summe"
    for para in tbl.cell(22, 0).paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.font.size = Pt(12)

    # -------------------------
    # Zeile 24: "Übertrag" + Merges wie Header (2&3, 4&5, 6&7, 8&9)
    # -------------------------
    tbl.cell(23, 0).text = "Übertrag"
    for para in tbl.cell(23, 0).paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.font.size = Pt(12)

    for left in (1, 3, 5, 7):
        merge_pair(23, left)

    # In den restlichen Zellen bleibt alles leer (so wie du es möchtest).


def _add_page_break(doc) -> None:
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# -----------------------------------------------------------------------------
# Route: Sammel-DOCX (Variante B, alles per Code)
# -----------------------------------------------------------------------------
@bp.get("/tournaments/<int:tournament_id>/rounds/<int:round_no>/tablesheets-docx-merged")
def tournament_round_tablesheets_docx_merged(tournament_id: int, round_no: int):
    try:
        from docx import Document  # type: ignore
    except ModuleNotFoundError:
        flash("DOCX-Export nicht verfügbar: Paket 'python-docx' ist nicht installiert.", "error")
        return redirect(url_for("tournaments.tournament_round_view", tournament_id=tournament_id, round_no=round_no))

    with db.connect() as con:
        t = _get_tournament(con, tournament_id)
        if not t:
            flash("Turnier nicht gefunden.", "error")
            return redirect(url_for("tournaments.tournaments_list"))

        tables = _fetch_round_tables(con, tournament_id, round_no)
        if not tables:
            flash(f"Keine vollständigen 4er-Tische für Runde {round_no} gefunden (Auslosung fehlt?).", "error")
            return redirect(url_for("tournaments.tournament_round_view", tournament_id=tournament_id, round_no=round_no))

        title = str(t["title"] or "").strip() or "Turnier"

    # -------------------------
    # Dokument bauen
    # -------------------------
    doc = Document()

    # Sicherstellen: keine Kopf-/Fußzeilen-Inhalte (normalerweise leer, aber wir leeren defensiv)
    for sec in doc.sections:
        sec.header.is_linked_to_previous = False
        sec.footer.is_linked_to_previous = False
        for p in list(sec.header.paragraphs):
            p.text = ""
        for p in list(sec.footer.paragraphs):
            p.text = ""

    # pro Tisch: Seite 1 (1-20), Seitenumbruch, Seite 2 (21-40), Seitenumbruch (außer am Ende optional)
    for idx, table in enumerate(tables):
        # Seite 1
        _build_one_page(
            doc,
            tournament_title=title,
            round_no=round_no,
            table_no=table.table_no,
            seats=table.seats,
            game_start_no=1,
        )
        _add_page_break(doc)

        # Seite 2 (21..40)
        _build_one_page(
            doc,
            tournament_title=title,
            round_no=round_no,
            table_no=table.table_no,
            seats=table.seats,
            game_start_no=21,
        )

        # Nach Seite 2: Seitenumbruch für nächsten Tisch (aber nicht zwingend am Dokumentende)
        if idx < (len(tables) - 1):
            _add_page_break(doc)

    # -------------------------
    # Ausliefern
    # -------------------------
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)

    fn = f"tischblaetter_t{int(tournament_id)}_r{int(round_no)}.docx"
    return send_file(
        mem,
        as_attachment=True,
        download_name=fn,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        max_age=0,
    )