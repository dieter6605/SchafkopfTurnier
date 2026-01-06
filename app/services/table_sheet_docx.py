# app/services/table_sheet_docx.py
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict

from docx import Document


def render_table_sheet_docx(*, template_path: str, ctx: Dict[str, Any]) -> bytes:
    """
    Rendert aus einer DOCX-Vorlage ein Tischblatt (2 Seiten).
    Ersetzt Platzhalter wie {{TURNIER_NAME}}, {{RUNDE_NR}}, {{TISCH_NR}},
    {{P1_NO}}, {{P1_NAME}}, {{P1_EMAIL}}, ... bis P4.
    """
    doc = Document(template_path)

    def repl(text: str) -> str:
        out = text
        for k, v in ctx.items():
            out = out.replace(f"{{{{{k}}}}}", "" if v is None else str(v))
        return out

    # Ersetzen in allen Absätzen (inkl. Header)
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.runs:
                for r in p.runs:
                    r.text = repl(r.text)
            else:
                p.text = repl(p.text)

    for p in doc.paragraphs:
        if p.runs:
            for r in p.runs:
                r.text = repl(r.text)
        else:
            p.text = repl(p.text)

    # Ersetzen in Tabellen (Zellen/Absätze)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.runs:
                        for r in p.runs:
                            r.text = repl(r.text)
                    else:
                        p.text = repl(p.text)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()